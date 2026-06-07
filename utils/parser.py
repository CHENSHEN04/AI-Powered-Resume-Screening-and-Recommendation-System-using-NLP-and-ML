"""
Resume Parser Module
====================
Extracts text and metadata from PDF and DOCX resume files.

Implements the parsing pipeline described in OUTPUT_SPECIFICATION.md section 2.1.
"""

from dataclasses import dataclass
from typing import Optional
from io import BytesIO
from pathlib import Path

from utils.errors import ErrorCode, AppError, get_error, log_error


# ==============================================================================
# Data Classes
# ==============================================================================

@dataclass
class ParseResult:
    """Result of parsing a resume file."""
    success: bool
    text: str
    page_count: int = 0
    confidence: float = 0.0
    error_code: Optional[str] = None
    error: Optional[AppError] = None
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


# ==============================================================================
# Resume Parser Class
# ==============================================================================

class ResumeParser:
    """
    Parse PDF and DOCX resumes to extract text content.
    
    Supports:
    - PDF files (via pdfplumber and PyMuPDF fallback)
    - DOCX files (via python-docx)
    - Error handling for corrupted/password-protected files
    """
    
    def __init__(self):
        """Initialize the parser."""
        pass
    
    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        """
        Parse a resume file and extract text.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to determine type)
        
        Returns:
            ParseResult with extracted text and metadata
        """
        # 1. Validate File
        is_valid, error = self._validate_file(file_bytes, filename)
        if not is_valid:
            return ParseResult(
                success=False,
                text="",
                error_code=error.code.value,
                error=error
            )

        ext = Path(filename).suffix.lower()
        
        result = None
        if ext == ".pdf":
            result = self.parse_pdf(file_bytes)
        elif ext == ".docx":
            result = self.parse_docx(file_bytes)
        else:
            # Should be caught by validate_file, but safety net
            error = get_error(ErrorCode.INVALID_FILE_TYPE)
            return ParseResult(
                success=False,
                text="",
                error_code=error.code.value,
                error=error
            )
            
        # 2. Check Page Limit (Warning)
        if result.success and result.page_count > 3:
            result.metadata["warning"] = "Page count exceeds recommended limit (3 pages)."
            
        return result

    def _validate_file(self, file_bytes: bytes, filename: str) -> tuple[bool, Optional[AppError]]:
        """
        Validate file size, type (MIME), and content.
        """
        # Check Empty
        if len(file_bytes) == 0:
            return False, get_error(ErrorCode.EMPTY_FILE)
            
        # Check Size (10MB limit)
        if len(file_bytes) > 10 * 1024 * 1024:
            return False, get_error(ErrorCode.FILE_TOO_LARGE)

        # Check MIME Type
        try:
            import magic
            mime = magic.Magic(mime=True)
            file_type = mime.from_buffer(file_bytes)
            
            allowed_mimes = [
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", # .docx
                "application/msword" # .doc (legacy, maybe treat as invalid or try docx parser)
            ]
            
            # Relaxed check for docx which can sometimes appear as zip
            if file_type not in allowed_mimes and file_type != "application/zip":
                return False, get_error(ErrorCode.INVALID_FILE_TYPE)
                
        except ImportError:
            # Fallback if python-magic not installed
            pass
        except Exception as e:
            # Log but don't block if magic fails
            pass

        return True, None
    
    def parse_pdf(self, file_bytes: bytes) -> ParseResult:
        """
        Parse a PDF file using pdfplumber with PyMuPDF fallback.
        
        Args:
            file_bytes: PDF file content as bytes
        
        Returns:
            ParseResult with extracted text
        """
        import pdfplumber
        import fitz  # PyMuPDF
        from io import BytesIO
        
        pdf_file = BytesIO(file_bytes)
        text_chunks = []
        page_count = 0
        
        try:
            # Try pdfplumber first (better for structured text)
            with pdfplumber.open(pdf_file) as pdf:
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_chunks.append(page_text)
            
            full_text = "\n\n".join(text_chunks)
            
            # Check if we got any text
            if not full_text.strip():
                # Try PyMuPDF as fallback
                pdf_file.seek(0)
                doc = fitz.open(stream=pdf_file, filetype="pdf")
                page_count = len(doc)
                
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text_chunks.append(page_text)
                
                doc.close()
                full_text = "\n\n".join(text_chunks)
            
            # If still no text, return error
            if not full_text.strip():
                error = get_error(ErrorCode.NO_TEXT_EXTRACTED)
                log_error(error, {"file_type": "pdf", "page_count": page_count})
                return ParseResult(
                    success=False,
                    text="",
                    page_count=page_count,
                    error_code=error.code.value,
                    error=error
                )
            
            # Calculate confidence based on text length and structure
            confidence = self._calculate_confidence(full_text, page_count)
            
            return ParseResult(
                success=True,
                text=full_text,
                page_count=page_count,
                confidence=confidence,
                metadata={
                    "file_type": "pdf",
                    "text_length": len(full_text),
                    "parser": "pdfplumber"
                }
            )
        
        except Exception as e:
            # Check for password protection
            error_str = str(e).lower()
            if "password" in error_str or "encrypted" in error_str:
                error = get_error(ErrorCode.PASSWORD_PROTECTED)
            else:
                error = get_error(ErrorCode.PDF_EXTRACTION_FAILED)
                error.message = f"PDF extraction failed: {str(e)}"
            
            log_error(error, {"exception": str(e), "file_type": "pdf"})
            return ParseResult(
                success=False,
                text="",
                page_count=page_count,
                error_code=error.code.value,
                error=error
            )
    
    def parse_docx(self, file_bytes: bytes) -> ParseResult:
        """
        Parse a DOCX file using python-docx.
        
        Args:
            file_bytes: DOCX file content as bytes
        
        Returns:
            ParseResult with extracted text
        """
        from docx import Document
        from io import BytesIO
        
        try:
            docx_file = BytesIO(file_bytes)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text_chunks = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_chunks.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_chunks.append(row_text)
            
            full_text = "\n".join(text_chunks)
            
            # Check if we got any text
            if not full_text.strip():
                error = get_error(ErrorCode.NO_TEXT_EXTRACTED)
                log_error(error, {"file_type": "docx"})
                return ParseResult(
                    success=False,
                    text="",
                    error_code=error.code.value,
                    error=error
                )
            
            # Estimate page count (rough estimate: 500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)
            
            # Calculate confidence
            confidence = self._calculate_confidence(full_text, page_count)
            
            return ParseResult(
                success=True,
                text=full_text,
                page_count=page_count,
                confidence=confidence,
                metadata={
                    "file_type": "docx",
                    "text_length": len(full_text),
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables)
                }
            )
        
        except Exception as e:
            error = get_error(ErrorCode.DOCX_EXTRACTION_FAILED)
            error.message = f"DOCX extraction failed: {str(e)}"
            log_error(error, {"exception": str(e), "file_type": "docx"})
            return ParseResult(
                success=False,
                text="",
                error_code=error.code.value,
                error=error
            )
    
    def _calculate_confidence(self, text: str, page_count: int) -> float:
        """
        Calculate parsing confidence based on text characteristics.
        
        Args:
            text: Extracted text
            page_count: Number of pages
        
        Returns:
            Confidence score (0.0 to 1.0)
        """
        confidence = 0.5  # Base confidence
        
        # Boost confidence if we have reasonable text length
        text_length = len(text.strip())
        if text_length > 200:
            confidence += 0.2
        if text_length > 500:
            confidence += 0.1
        
        # Boost confidence if we detect common resume sections
        resume_keywords = [
            "experience", "education", "skills", "projects",
            "work", "university", "college", "degree"
        ]
        text_lower = text.lower()
        keyword_count = sum(1 for kw in resume_keywords if kw in text_lower)
        confidence += min(0.2, keyword_count * 0.05)
        
        # Cap confidence at 1.0
        return min(1.0, confidence)

    def convert_pdf_to_image(self, file_bytes: bytes) -> Optional[bytes]:
        """
        Render the first page of the PDF to a PNG byte array in memory.
        """
        import fitz
        from io import BytesIO
        
        try:
            doc = fitz.open(stream=BytesIO(file_bytes), filetype="pdf")
            if len(doc) == 0:
                return None
            page = doc[0]  # Render first page
            pix = page.get_pixmap(dpi=150)  # High-quality rendering
            png_bytes = pix.tobytes("png")
            doc.close()
            return png_bytes
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"Failed to render PDF to PNG: {e}")
            return None

    def extract_font_metadata(self, file_bytes: bytes) -> list:
        """
        Extract detailed text formatting, fonts, sizes and positions.
        """
        import fitz
        from io import BytesIO
        
        metadata = []
        try:
            doc = fitz.open(stream=BytesIO(file_bytes), filetype="pdf")
            if len(doc) > 0:
                page = doc[0]  # First page
                
                # Append page_info element containing dimensions, safely structured
                # to not cause KeyErrors in existing code/tests.
                metadata.append({
                    "text": "",
                    "font": "",
                    "size": 0.0,
                    "bbox": [0.0, 0.0, 0.0, 0.0],
                    "type": "page_info",
                    "width": round(page.rect.width, 1),
                    "height": round(page.rect.height, 1)
                })
                
                blocks = page.get_text("dict")["blocks"]
                for b in blocks:
                    if b.get("type") == 0:  # Text block
                        for line in b["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if len(text) > 2:  # Ignore trivial characters
                                    metadata.append({
                                        "text": text,
                                        "font": span["font"],
                                        "size": round(span["size"], 1),
                                        "bbox": [round(x, 1) for x in span["bbox"]]
                                    })
            doc.close()
        except Exception:
            pass
        return metadata
